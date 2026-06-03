"""Text extraction from business documents."""

import re
from pathlib import Path
from typing import List, Optional


# PDF extraction
def extract_pdf_as_markdown(file_path: str) -> str:
    """Extract PDF as Markdown using PyMuPDF4LLM."""
    try:
        import pymupdf4llm
        return pymupdf4llm.to_markdown(file_path)
    except ImportError:
        import pymupdf
        doc = pymupdf.open(file_path)
        parts = []
        for page in doc:
            text = page.get_text("text")
            if text.strip():
                parts.append(f"--- Trang {page.number + 1} ---\n{text}")
        doc.close()
        return "\n\n".join(parts)


# DOCX extraction
def extract_docx_text(file_path: str) -> str:
    """Extract text + tables from DOCX."""
    from docx import Document
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            if para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                prefix = "#" * int(level) if level.isdigit() else "##"
                parts.append(f"{prefix} {para.text}")
            else:
                parts.append(para.text)
    for i, table in enumerate(doc.tables):
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        rows = []
        for row in table.rows[1:]:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):
                rows.append(dict(zip(headers, row_data)))
        if rows:
            header_line = "| " + " | ".join(headers) + " |"
            separator = "| " + " | ".join(["---"] * len(headers)) + " |"
            data_lines = ["| " + " | ".join(str(row.get(h, "")) for h in headers) + " |" for row in rows]
            parts.append(f"\n**Bảng {i+1}:**\n" + "\n".join([header_line, separator] + data_lines))
    return "\n\n".join(parts)


# Excel extraction
def extract_excel_text(file_path: str) -> str:
    """Extract text from Excel as markdown tables."""
    import openpyxl
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(cell) if cell is not None else f"C{i+1}" for i, cell in enumerate(rows[0])]
        parts.append(f"## Sheet: {sheet_name}")
        data_rows = rows[1:201]  # Limit 200 rows
        if data_rows:
            header_line = "| " + " | ".join(headers) + " |"
            separator = "| " + " | ".join(["---"] * len(headers)) + " |"
            body_lines = ["| " + " | ".join(str(cell) if cell is not None else "" for cell in row) + " |" for row in data_rows]
            parts.append("\n".join([header_line, separator] + body_lines))
            if len(rows) > 201:
                parts.append(f"... (tổng {len(rows)} dòng)")
    wb.close()
    return "\n\n".join(parts)


# CSV extraction
def extract_csv_text(file_path: str) -> str:
    """Extract text from CSV as markdown table."""
    import pandas as pd
    df = pd.read_csv(file_path, nrows=200)
    headers = list(df.columns)
    header_line = "| " + " | ".join(headers) + " |"
    separator = "| " + " | ".join(["---"] * len(headers)) + " |"
    body_lines = ["| " + " | ".join(str(v) if pd.notna(v) else "" for v in row) + " |" for _, row in df.iterrows()]
    result = "\n".join([header_line, separator] + body_lines)
    total_rows = len(pd.read_csv(file_path))
    if total_rows > 200:
        result += f"\n... (tổng {total_rows} dòng)"
    return result


# Unified dispatcher
def extract_text(file_path: str) -> str:
    """Extract text from any supported file type."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_pdf_as_markdown(file_path)
    elif ext == ".docx":
        return extract_docx_text(file_path)
    elif ext in (".xlsx", ".xls"):
        return extract_excel_text(file_path)
    elif ext == ".csv":
        return extract_csv_text(file_path)
    else:
        raise ValueError(f"Loại file không hỗ trợ: {ext}")
